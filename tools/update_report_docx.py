import os
import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]


def read_text(rel_path: str) -> str:
    return (ROOT / rel_path).read_text(encoding="utf-8")


def parse_catalog_table(md: str) -> dict[str, dict[str, str]]:
    """
    Parse markdown tables from VULNERABILITY_CATALOG.md.
    Returns mapping: id -> {category, level, support, what, verify}
    """
    rows: dict[str, dict[str, str]] = {}
    for line in md.splitlines():
        line = line.strip()
        if not line.startswith("|"):
            continue
        # Skip header separators
        if re.match(r"^\|\s*-+\s*\|", line):
            continue
        parts = [p.strip() for p in line.strip("|").split("|")]
        if len(parts) < 6:
            continue
        vuln_id = parts[0].strip()
        if not (vuln_id.startswith("`") and vuln_id.endswith("`")):
            continue
        vuln_id = vuln_id.strip("`")
        rows[vuln_id] = {
            "category": parts[1],
            "level": parts[2],
            "support": parts[3],
            "what": parts[4],
            "verify": parts[5],
        }
    return rows


def parse_scenarios(md: str) -> dict[str, dict[str, str]]:
    """
    Parse VULNERABILITY_SCENARIOS.md sections:
      ### `id`
      - **Что делает**: ...
      - **Учебный сценарий**: ...
      - **Проверка**: ...
    Returns mapping: id -> {what, scenario, verify}
    """
    out: dict[str, dict[str, str]] = {}
    current_id: str | None = None
    buf: list[str] = []

    def flush():
        nonlocal current_id, buf
        if not current_id:
            return
        text = "\n".join(buf)

        def grab(label: str) -> str:
            # capture from "- **Label**:" to next "- **" or end
            m = re.search(
                rf"-\s+\*\*{re.escape(label)}\*\*:\s*(.*?)(?=\n-\s+\*\*|\Z)",
                text,
                flags=re.S,
            )
            if not m:
                return ""
            return re.sub(r"\n\s*- ", "\n- ", m.group(1).strip())

        out[current_id] = {
            "what": grab("Что делает"),
            "scenario": grab("Учебный сценарий"),
            "verify": grab("Проверка"),
        }
        current_id = None
        buf = []

    for line in md.splitlines():
        m = re.match(r"^###\s+`([^`]+)`\s*$", line.strip())
        if m:
            flush()
            current_id = m.group(1).strip()
            buf = []
            continue
        if current_id is not None:
            buf.append(line.rstrip())
    flush()
    return out


def add_bullets(doc: Document, text: str):
    if not text:
        return
    # Keep simple: lines beginning with "- " become bullet points.
    lines = [ln.rstrip() for ln in text.splitlines() if ln.strip()]
    for ln in lines:
        if ln.lstrip().startswith("- "):
            doc.add_paragraph(ln.lstrip()[2:], style="List Bullet")
        else:
            doc.add_paragraph(ln)


@dataclass
class Counters:
    figures: int = 0
    screenshots: int = 0
    listings: int = 0
    tables: int = 0


def add_placeholder(doc: Document, kind: str, title: str, c: Counters):
    kind = kind.lower()
    if kind == "figure":
        c.figures += 1
        label = f"Рисунок {c.figures}"
        doc.add_paragraph(f"[МЕСТО ДЛЯ СХЕМЫ/РИСУНКА: {label} — {title}]")
        return
    if kind == "screenshot":
        c.screenshots += 1
        label = f"Скриншот {c.screenshots}"
        doc.add_paragraph(f"[МЕСТО ДЛЯ СКРИНШОТА: {label} — {title}]")
        return
    if kind == "listing":
        c.listings += 1
        label = f"Листинг {c.listings}"
        doc.add_paragraph(f"[МЕСТО ДЛЯ ЛИСТИНГА: {label} — {title}]")
        return
    if kind == "table":
        c.tables += 1
        label = f"Таблица {c.tables}"
        doc.add_paragraph(f"[МЕСТО ДЛЯ ТАБЛИЦЫ: {label} — {title}]")
        return
    doc.add_paragraph(f"[МЕСТО ДЛЯ ВСТАВКИ: {title}]")


def add_code_block(doc: Document, code: str):
    """
    Minimal monospace block (without relying on custom styles).
    """
    p = doc.add_paragraph()
    for i, line in enumerate(code.rstrip("\n").splitlines()):
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        if i != len(code.rstrip("\n").splitlines()) - 1:
            p.add_run("\n")


def extract_fix_actions(scenario_text: str) -> str:
    """
    Best-effort: extract 'исправить:' fragments from scenario bullets.
    """
    if not scenario_text:
        return ""
    lines = [ln.strip() for ln in scenario_text.splitlines() if ln.strip()]
    fixes = []
    for ln in lines:
        if "исправ" in ln.lower():
            fixes.append(ln.lstrip("- ").strip())
    return "\n".join(f"- {f}" for f in fixes[:10])


def category_rationale(category: str, scope: str) -> str:
    """
    category: string number from catalog (e.g. '1')
    scope: linux|windows|ad
    """
    c = (category or "").strip()
    if c == "1":
        return (
            "Группа «Аутентификация и парольные политики» выбрана, потому что компрометация учётных записей остаётся одним из "
            "самых частых путей входа в инфраструктуру. Слабая password policy, отсутствие lockout, разрешение парольного входа по SSH "
            "или наличие «слабых» админ‑паролей делают возможными перебор и использование утёкших учетных данных."
        )
    if c == "2":
        return (
            "Группа «Привилегии и контроль доступа» выбрана из‑за критичности принципа наименьших привилегий. "
            "Лишние администраторы, отключённый UAC, SUID/SGID backdoor или NOPASSWD в sudo резко упрощают повышение привилегий "
            "и закрепление в системе."
        )
    if c == "3":
        return (
            "Группа «Удалённый доступ» отражает риски, связанные с RDP/SSH/WinRM: расширение поверхности атаки, "
            "возможность brute-force, lateral movement и перехват учетных данных. В доменной среде важно ограничивать, "
            "кто и откуда может подключаться, и обеспечить безопасные параметры (NLA, MFA, key-based auth и т.д.)."
        )
    if c == "4":
        return (
            "Группа «Файлы и права доступа» выбрана, потому что ошибки прав (world-writable, небезопасный PATH, "
            "хранение секретов в доступных файлах) приводят к утечкам данных и локальной эскалации привилегий."
        )
    if c == "5":
        return (
            "Группа «Обновления и управление патчами» отражает распространённую проблему: отключение обновлений ради «стабильности». "
            "Это приводит к накоплению известных уязвимостей и увеличивает вероятность успешной эксплуатации."
        )
    if c == "6":
        return (
            "Группа «Сервисы и сетевые протоколы» включает небезопасные или устаревшие протоколы (SMB1, NTLMv1, Telnet/FTP, RemoteRegistry). "
            "Такие настройки повышают риск MITM/relay атак, утечек учетных данных и удалённого выполнения кода."
        )
    if c == "7":
        return (
            "Группа «Сетевой периметр и Firewall» выбрана, потому что отключённый или неправильно настроенный firewall "
            "делает хост доступным для сканирования и эксплуатации из соседних сегментов."
        )
    if c == "8":
        return (
            "Группа «Логирование и аудит» критична для расследования и обнаружения атак. "
            "Выключенный аудит, низкое удержание журналов, отключённое PowerShell логирование ухудшают видимость и облегчают сокрытие следов."
        )
    if c == "9":
        return (
            "Группа «Антивирус и базовые средства защиты» покрывает ситуации, когда защита отключается или ослабляется. "
            "В учебной среде это демонстрирует, как быстро растёт риск при отсутствии базовых средств защиты и мониторинга."
        )
    # backward-compat (в старых черновиках/таблицах встречалось 10)
    if c == "10":
        return category_rationale("9", scope)
    # generic
    base = (
        "Данная категория включена в учебный каталог, так как отражает типовой класс ошибок администрирования, "
        "которые приводят к повышенному риску компрометации."
    )
    if scope == "ad":
        base += " В доменной среде последствия ошибок усиливаются из‑за централизованного распространения политик (GPO)."
    return base


def evidence_template(scope: str) -> str:
    if scope == "linux":
        return (
            "- вывод команд проверки (до/после);\n"
            "- фрагменты конфигураций (/etc/ssh/sshd_config, /etc/sudoers.d/*, /etc/pam.d/*);\n"
            "- статус служб (systemctl status ...);\n"
            "- скриншот консоли Proxmox (если нужно) или лог выполнения Ansible."
        )
    if scope == "windows":
        return (
            "- вывод PowerShell/CLI команд (до/после): net accounts, Get-ItemProperty, Get-Service, Get-NetFirewallProfile;\n"
            "- подтверждение политики (gpresult /r, RSOP) при применении GPO;\n"
            "- скриншоты ключевых экранов (Local Security Policy / Group Policy) при необходимости."
        )
    return (
        "- вывод команд DC/AD: Get-ADUser/Get-ADGroupMember/Get-ADDefaultDomainPasswordPolicy;\n"
        "- подтверждение наличия/линковки GPO и области применения;\n"
        "- скриншоты GPMC (Group Policy Management) при необходимости."
    )


def category_name(category: str) -> str:
    c = (category or "").strip()
    names = {
        "1": "Аутентификация и парольные политики",
        "2": "Привилегии и контроль доступа",
        "3": "Удалённый доступ",
        "4": "Файлы и права доступа",
        "5": "Обновления и управление патчами",
        "6": "Сервисы и сетевые протоколы",
        "7": "Сетевой периметр и Firewall",
        "8": "Логирование и аудит",
        "9": "Антивирус и базовые средства защиты",
        # backward-compat (в старых черновиках/таблицах встречалось 10)
        "10": "Антивирус и базовые средства защиты",
    }
    return names.get(c, "Прочее")


def estimate_pages(word_count: int) -> int:
    # Rough heuristic for Word (A4, 12pt, 1.5 spacing) varies a lot.
    # We use conservative 320 words/page to avoid underestimating.
    return max(1, round(word_count / 320))


def count_words_doc(doc: Document) -> int:
    total = 0
    for p in doc.paragraphs:
        total += len(re.findall(r"[A-Za-zА-Яа-я0-9]+", p.text))
    return total


def add_markdown_as_text(doc: Document, md: str, counters: Counters):
    """
    Very lightweight Markdown → docx text:
    - '#', '##', '###' become headings
    - '- ' become bullets
    - fenced code blocks become listing placeholders
    This is intentionally simple and deterministic for a thesis appendix.
    """
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    for raw in md.splitlines():
        line = raw.rstrip("\n")
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line.strip().strip("`").strip()
                code_lines = []
            else:
                # close code block
                snippet = "\n".join(code_lines[:60]).rstrip()
                title = f"Фрагмент кода ({code_lang or 'code'})"
                add_placeholder(doc, "listing", title, counters)
                if snippet:
                    add_code_block(doc, snippet)
                in_code = False
                code_lang = ""
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        s = line.strip()
        if not s:
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:], level=3)
            continue
        if s.startswith("## "):
            doc.add_heading(s[3:], level=2)
            continue
        if s.startswith("# "):
            doc.add_heading(s[2:], level=1)
            continue
        if s.startswith("- "):
            doc.add_paragraph(s[2:], style="List Bullet")
            continue
        # keep plain lines as paragraphs
        doc.add_paragraph(s)


def save_docx_with_fallback(doc: Document, target_path: Path) -> Path:
    """
    Word on Windows often locks an opened .docx, causing PermissionError.
    In that case we write to a sibling file with suffix '__generated'.
    """
    try:
        doc.save(target_path)
        return target_path
    except PermissionError:
        base = target_path
        # Make unique fallback name
        for n in range(1, 50):
            suffix = "__generated" if n == 1 else f"__generated_{n}"
            alt = base.with_name(f"{base.stem}{suffix}{base.suffix}")
            try:
                doc.save(alt)
                return alt
            except PermissionError:
                continue
        # Last resort: ASCII filename
        alt = base.with_name("report__generated.docx")
        doc.save(alt)
        return alt


def add_example_vuln_impl(doc: Document, title: str, rel_task_path: str, counters: Counters, max_lines: int = 60):
    """
    Add a short example implementation excerpt for a single vulnerability task file.
    """
    doc.add_heading(title, level=3)
    add_placeholder(doc, "listing", f"Пример реализации: {rel_task_path}", counters)
    try:
        raw = read_text(rel_task_path)
        lines = raw.splitlines()
        snippet = "\n".join(lines[:max_lines]).rstrip()
        if snippet:
            add_code_block(doc, snippet)
    except Exception:
        doc.add_paragraph("[Не удалось автоматически прочитать файл реализации примера. Проверьте путь в репозитории.]")


def main():
    # Find target docx file in repo root.
    # Prefer the "main" report file (non-generated), because __generated is a fallback when Word locks the main file.
    docx_files = [f for f in os.listdir(ROOT) if f.lower().endswith(".docx")]
    docx_files = [f for f in docx_files if not f.startswith("~$")]  # ignore MS Word temp files
    if not docx_files:
        raise SystemExit("No .docx files found in repo root.")

    def is_generated(name: str) -> bool:
        n = name.lower()
        return "__generated" in n or "report__generated" in n

    non_generated = [f for f in docx_files if not is_generated(f)]
    if non_generated:
        # Prefer the shortest name (usually the base file) and newest mtime as tie-breaker.
        non_generated.sort(key=lambda n: (len(n), -os.path.getmtime(ROOT / n)))
        target_name = non_generated[0]
    else:
        # If only generated variants exist, take the newest one.
        docx_files.sort(key=lambda n: -os.path.getmtime(ROOT / n))
        target_name = docx_files[0]

    target_path = ROOT / target_name

    prereq = read_text("docs/guides/PREREQUISITES.md")
    training = read_text("docs/guides/TRAINING_SCENARIO.md")
    catalog_md = read_text("docs/guides/VULNERABILITY_CATALOG.md")
    scenarios_md = read_text("docs/guides/VULNERABILITY_SCENARIOS.md")

    catalog = parse_catalog_table(catalog_md)
    scen = parse_scenarios(scenarios_md)
    counters = Counters()

    # Sort IDs by prefix for readability
    def sort_key(v: str):
        return (v.split(".", 1)[0], v)

    linux_ids = sorted([k for k in catalog if k.startswith("linux.")], key=sort_key)
    win_ids = sorted([k for k in catalog if k.startswith("windows.")], key=sort_key)
    ad_ids = sorted([k for k in catalog if k.startswith("ad.")], key=sort_key)

    doc = Document()

    # Титульный лист (заготовка)
    doc.add_heading("ОТЧЁТ ПО ВЫПУСКНОЙ КВАЛИФИКАЦИОННОЙ РАБОТЕ", 0)
    doc.add_paragraph("Тема: Разработка платформы киберучений по защите операционных систем на базе Proxmox VE (Terraform + Ansible)")
    doc.add_paragraph("[МЕСТО ДЛЯ ТИТУЛЬНЫХ ДАННЫХ ПО ШАБЛОНУ ВУЗа: ВУЗ/кафедра/направление/группа/ФИО/руководитель/город/год]")
    doc.add_paragraph(f"Дата автоматической сборки текста: {date.today().isoformat()}")
    doc.add_page_break()

    # Аннотация
    doc.add_heading("Аннотация", level=1)
    doc.add_paragraph(
        "В работе описана платформа киберучений по курсу «Защита ОС», построенная на принципах Infrastructure as Code. "
        "Платформа обеспечивает автоматизированное развертывание учебных стендов в Proxmox VE с помощью Terraform, "
        "а также пост-настройку операционных систем и доменной инфраструктуры (Active Directory/GPO) с помощью Ansible. "
        "В стендах реализован каталог типовых уязвимостей/ошибок администрирования, которые включаются по профилям, "
        "что позволяет формировать разные варианты лабораторных работ и практикумов."
    )
    doc.add_paragraph(
        "Ключевые слова: Proxmox VE, Terraform, Ansible, Active Directory, GPO, SSH, cloud-init, IaC, hardening, киберучения."
    )
    doc.add_page_break()

    # Содержание (заготовка)
    doc.add_heading("Содержание", level=1)
    doc.add_paragraph("[ОГЛАВЛЕНИЕ: в Word вставьте автоматически (Ссылки → Оглавление) после финального форматирования]")
    doc.add_page_break()

    # Введение
    doc.add_heading("Введение", level=1)
    doc.add_paragraph(
        "Обеспечение безопасности операционных систем является базовой задачей при эксплуатации любой информационной системы. "
        "На практике значительная часть инцидентов связана не с «неизвестными 0-day», а с типовыми ошибками администрирования: "
        "слабыми политиками паролей, отключенной регистрацией событий, избыточными правами, небезопасными сетевыми настройками и "
        "несвоевременными обновлениями. Отдельное место занимают доменные инфраструктуры на базе Active Directory, где ошибочная "
        "настройка групповых политик (GPO) может массово снижать уровень защищённости рабочих станций и серверов."
    )
    doc.add_paragraph(
        "Целью работы является разработка автоматизированной платформы киберучений по курсу «Защита ОС», позволяющей "
        "быстро разворачивать учебные стенды в среде виртуализации Proxmox VE и воспроизводить типовые уязвимости/ошибки конфигурации "
        "с последующей проверкой и устранением обучающимися."
    )
    doc.add_paragraph("Для достижения цели решались следующие задачи:")
    add_bullets(
        doc,
        "- определить типовые уязвимости и ошибки администрирования, характерные для небольших организаций (малые домены, ограниченные ресурсы и компетенции);\n"
        "- разработать учебные инфраструктуры, моделирующие конфигурационные ошибки и угрозы безопасности, встречающиеся в небольших организациях;\n"
        "- реализовать автоматизированное развертывание разработанных инфраструктур с использованием принципов IaC (Terraform + Ansible) в Proxmox VE;\n"
        "- подготовить документацию, каталог уязвимостей и методические материалы для обучения;\n"
        "- обеспечить наличие операторского доступа администратора к ВМ (SSH/учётные записи) и переносимость стендов.",
    )
    doc.add_paragraph(
        "Объект исследования: процессы развертывания и сопровождения учебной инфраструктуры операционных систем. "
        "Предмет исследования: методы автоматизации развертывания стендов и воспроизведения типовых уязвимостей "
        "в ОС и доменной инфраструктуре Active Directory."
    )
    doc.add_paragraph(
        "Практическая значимость работы заключается в получении воспроизводимого стенда «одной командой», "
        "который можно переносить между инсталляциями Proxmox и использовать для проведения лабораторных работ и практикумов."
    )
    doc.add_paragraph(
        "Структура отчёта: в разделе 1 приведены типовые уязвимости/ошибки администрирования и обоснование выбора категорий; "
        "в разделе 2 описано проектирование учебной площадки (архитектура, структура репозитория, механизм уязвимостей и расширение каталога); "
        "в разделе 3 представлены инструкции для администратора по развертыванию, диагностике и сопровождению стендов; "
        "в разделе 4 описаны сценарии тренировки и обучения; приложения содержат выдержки из методических материалов и каталога."
    )
    add_placeholder(doc, "figure", "Общая архитектура платформы (Proxmox ↔ Terraform ↔ Ansible ↔ стенды)", counters)
    doc.add_page_break()

    # 1) Типовые уязвимости (ставим ПЕРЕД описанием учебной среды)
    doc.add_heading("1. Типовые уязвимости (описание + сценарии реализации)", level=1)
    doc.add_paragraph(
        "Уязвимости реализованы как модули с ID и включаются списками в vuln_profiles соответствующего стенда. "
        "Применение выполняется ролями: vuln_linux, vuln_windows, vuln_ad (AD/GPO)."
    )
    doc.add_paragraph(
        "Подход «уязвимость как модуль» позволяет включать и выключать наборы уязвимостей без изменения кода ролей: "
        "администратор редактирует только списки ID в group_vars. Это повышает повторяемость занятий и упрощает расширение каталога."
    )
    doc.add_paragraph(
        "В текущей версии платформы реализовано 52 уязвимости/ошибки администрирования: "
        "17 для Linux, 18 для Windows и 17 для Active Directory/GPO. "
        "Такой набор подобран так, чтобы охватить наиболее частые классы проблем, которые проверяются при аудитах и "
        "встречаются в типовых атаках: слабые политики паролей и lockout, избыточные привилегии, небезопасные удалённые протоколы, "
        "отключенная защита/обновления, ослабленное логирование/аудит, и доменные политики (GPO), распространяющие ошибочную конфигурацию на множество хостов."
    )
    doc.add_heading("1.1 Группы типовых уязвимостей и принципы отбора", level=2)
    doc.add_paragraph(
        "В проекте уязвимости группируются по смысловым разделам (аутентификация, привилегии, удалённый доступ и т.д.). "
        "Важно: в коде автоматизации (Ansible/Terraform) эти разделы **не используются** — автоматизация оперирует только **ID уязвимостей**. "
        "Числовые значения (1..9) встречаются только в документации (каталог уязвимостей) как авторская классификация для удобства навигации."
    )
    add_bullets(
        doc,
        "- выбранные группы в текущей версии:\n"
        f"  - {category_name('1')}\n"
        f"  - {category_name('2')}\n"
        f"  - {category_name('3')}\n"
        f"  - {category_name('4')}\n"
        f"  - {category_name('5')}\n"
        f"  - {category_name('6')}\n"
        f"  - {category_name('7')}\n"
        f"  - {category_name('8')}\n"
        f"  - {category_name('9')}\n",
    )
    add_bullets(
        doc,
        "- фокус на типовых ошибках администрирования, а не на редких уязвимостях приложений;\n"
        "- воспроизводимость в лабораторной среде (уязвимость можно включить/выключить и проверить «до/после»);\n"
        "- безопасность учебной эксплуатации: сценарии рассчитаны на изолированную сеть;\n"
        "- привязка к практическим техникам атак (брутфорс, lateral movement, relay, повышение привилегий, сокрытие следов);\n"
        "- покрытие OS-level и доменных политик (AD/GPO), так как в реальности домен — главный механизм централизованной конфигурации.",
    )
    doc.add_heading("1.2 Соотнесение с моделью атак (MITRE ATT&CK)", level=2)
    doc.add_paragraph(
        "Для обоснования реализованных уязвимостей используется модель MITRE ATT&CK: "
        "каждая уязвимость упрощает одну или несколько техник злоумышленника (Credential Access, Lateral Movement, Defense Evasion и др.). "
        "В учебном процессе студент должен связать найденную конфигурацию с возможным вектором атаки и предложить корректное исправление."
    )
    add_placeholder(doc, "table", "Матрица: уязвимости ↔ техники MITRE ATT&CK ↔ артефакты обнаружения ↔ исправления", counters)

    # Вместо перечисления всех уязвимостей с листингами в основной части:
    # - описываем механизм реализации и управления;
    # - полный перечень оставляем в приложении (каталог/сценарии);
    # - приводим несколько показательных примеров.
    doc.add_heading("1.3 Пояснение по выбранным уязвимостям (почему выбраны)", level=2)
    doc.add_paragraph(
        "Выбор уязвимостей в учебном каталоге делался не как «перечень эксплойтов», а как набор типовых ошибок администрирования, "
        "которые: (а) реально встречаются в инфраструктурах, (б) приводят к измеримому росту риска, (в) хорошо обучают базовым навыкам аудита и hardening."
    )
    doc.add_paragraph(
        "Основная идея — показать обучающемуся связь между конфигурацией и последствиями: "
        "от слабых политик аутентификации и избыточных прав — к захвату учёток и повышению привилегий; "
        "от небезопасных сетевых параметров — к lateral movement и relay‑атакам; "
        "от выключенных обновлений/защиты/аудита — к росту времени присутствия атакующего и усложнению расследования."
    )
    doc.add_paragraph("Почему выбраны именно следующие группы типовых уязвимостей:")
    # краткий блок «почему» по каждой категории (без перечисления всех уязвимостей)
    for cat in ["1", "2", "3", "4", "5", "6", "7", "8", "9"]:
        doc.add_paragraph(f"{category_name(cat)}: {category_rationale(cat, 'ad')}")
    doc.add_paragraph(
        "Отдельно выделены уязвимости Active Directory / GPO, потому что доменные политики способны массово распространять "
        "ошибочную конфигурацию на множество хостов, что приближает учебный стенд к реальной корпоративной среде."
    )
    add_placeholder(doc, "table", "Сводная таблица: категории → какие типовые ошибки покрываются → какие артефакты искать", counters)

    doc.add_heading("1.4 Где приведён полный каталог и сценарии", level=2)
    doc.add_paragraph(
        "Полный перечень уязвимостей (ID → категория → уровень → проверка) и сценарии заданий для обучающихся "
        "вынесены в методические материалы проекта и приведены в приложениях: "
        "`docs/guides/VULNERABILITY_CATALOG.md` и `docs/guides/VULNERABILITY_SCENARIOS.md`. "
        "Проектная реализация механизма уязвимостей (роли, профили, добавление/изменение) описана в главе 2."
    )
    doc.add_page_break()

    # 2) ПРОЕКТИРОВАНИЕ УЧЕБНОЙ ПЛОЩАДКИ
    doc.add_heading("2. ПРОЕКТИРОВАНИЕ УЧЕБНОЙ ПЛОЩАДКИ", level=1)
    doc.add_paragraph(
        "Раздел описывает ключевые проектные решения, архитектуру и структуру репозитория, а также механизм управления "
        "каталогом уязвимостей как набором модулей (включение по профилям, расширение и сопровождение)."
    )

    doc.add_heading("2.1 Требования и ключевые проектные решения", level=2)
    add_bullets(
        doc,
        "- воспроизводимость: стенд разворачивается одинаково при повторных запусках (idempotent-подход IaC);\n"
        "- переносимость: параметры Proxmox/сети (bridge, DHCP/статик, шлюз, DNS) задаются переменными;\n"
        "- управляемость: Linux и Windows администрируются по SSH, чтобы Ansible работал единообразно;\n"
        "- обязательный AD/DC: домен используется как базовый элемент курса; присоединение конкретной ВМ к домену выбирается настройкой;\n"
        "- разделение «уязвимости» и «секреты»: секретные значения не коммитятся (tfvars и пароли остаются локально);\n"
        "- расширяемость: уязвимости реализованы как отдельные модули Ansible с уникальными ID и единым шаблоном добавления.",
    )
    add_placeholder(doc, "figure", "Архитектурная схема: Proxmox ↔ (templates/Packer) ↔ Terraform ↔ Ansible ↔ стенды", counters)

    doc.add_heading("2.2 Архитектура и компоненты", level=2)
    add_bullets(
        doc,
        "- Proxmox VE: платформа виртуализации (KVM), где создаются/клонируются ВМ;\n"
        "- Terraform: создаёт ВМ из шаблонов, задаёт ресурсы, диски, сеть, cloud-init и включает QEMU Guest Agent;\n"
        "- Ansible: выполняет пост-настройку (учётные записи, (опц.) domain join, учебные уязвимости ОС и AD/GPO);\n"
        "- Packer/скрипты шаблонов: обеспечивают подготовку повторно используемых образов;\n"
        "- Документация: guides для администратора/преподавателя/студента (требования, сценарии, каталог уязвимостей).",
    )
    add_placeholder(doc, "figure", "Пайплайн: templates → terraform apply → ansible-playbook (DC → клиенты → DC)", counters)

    doc.add_heading("2.3 Структура репозитория (ключевые каталоги)", level=2)
    add_bullets(
        doc,
        "- stands/windows-stand: Windows стенд (DC + Windows Server + Windows 10) и инфраструктурные конфиги;\n"
        "- stands/linux-stand: Linux стенд (Linux Server + Linux WS) и инфраструктурные конфиги;\n"
        "- ansible/roles: переиспользуемые роли (учётные записи, AD join, AD org, уязвимости Linux/Windows/AD);\n"
        "- terraform/modules: общие Terraform-модули (при необходимости выноса повторяющихся блоков);\n"
        "- docs/guides: эксплуатационная документация и методические материалы;\n"
        "- tools: утилиты (в т.ч. генерация/обновление отчёта).",
    )
    add_placeholder(doc, "listing", "Дерево ключевых каталогов репозитория (stands/, ansible/roles/, docs/, tools/)", counters)

    doc.add_heading("2.4 Состав стендов и топология сети", level=2)
    doc.add_paragraph(
        "В репозитории реализованы два базовых стенда: Windows-стенд и Linux-стенд. "
        "Windows-стенд всегда включает Domain Controller (AD DS) и использует DC как DNS для домена. "
        "Linux-узлы могут (опционально) присоединяться к домену в зависимости от профиля join."
    )
    add_placeholder(doc, "figure", "Топология Windows-стенда: DC → Windows Server/Windows 10 (OU/GPO) + доступ обучающегося через Proxmox", counters)
    add_placeholder(doc, "figure", "Топология Linux-стенда: Linux Server + Linux WS (+ опциональный domain join к DC)", counters)

    doc.add_heading("2.5 Реализация уязвимостей как модулей (Ansible roles)", level=2)
    doc.add_paragraph(
        "Уязвимости реализованы как набор idempotent‑задач Ansible, сгруппированных в роли по типу цели: "
        "Linux (`vuln_linux`), Windows (`vuln_windows`) и Active Directory/GPO (`vuln_ad`). "
        "Каждая уязвимость имеет уникальный идентификатор (ID), который соответствует имени файла `tasks/vulns/<ID>.yml`."
    )
    add_bullets(
        doc,
        "- включение/выключение выполняется через `vuln_profiles.<host_group>` в `group_vars/all/vulnerabilities.yml`;\n"
        "- playbook передаёт выбранный список ID в роль как `vuln_enabled`;\n"
        "- роли содержат список допустимых ID (например `vuln_linux_known`, `vuln_windows_known`, `vuln_ad_known`) и валидируют входные данные;\n"
        "- часть задач реализована в режиме best-effort (не падают, если конкретный пакет/служба отсутствуют);\n"
        "- для Linux предусмотрены ветвления по семейству дистрибутива (Debian/RHEL) через `ansible_os_family`.",
    )
    add_placeholder(doc, "figure", "Схема: vuln_profiles → playbook → роль vuln_* → tasks/vulns/<id>.yml", counters)

    doc.add_heading("2.6 Управление профилями уязвимостей (изменение набора без правок кода)", level=2)
    doc.add_paragraph(
        "Администратор меняет учебное наполнение без правок кода ролей: достаточно отредактировать списки ID в `vuln_profiles` "
        "и повторно запустить соответствующий плейбук. Это позволяет иметь несколько «пакетов» уязвимостей под разные темы "
        "и уровни подготовки."
    )
    add_placeholder(doc, "listing", "Пример: vulnerabilities.yml (включение/выключение ID в профиле)", counters)
    add_bullets(
        doc,
        "- включить: добавить ID в список `vuln_profiles.<group>`;\n"
        "- выключить: удалить ID из списка;\n"
        "- применить изменения: повторно запустить `ansible-playbook ...` для нужного хоста/группы;\n"
        "- проверить: выполнить команды из раздела «Как проверить» в каталоге уязвимостей.",
    )

    doc.add_heading("2.7 Добавление и изменение уязвимостей (расширение каталога)", level=2)
    doc.add_paragraph(
        "Каталог расширяется по единому шаблону. Это важно для сопровождения: новые уязвимости добавляются предсказуемо, "
        "в одинаковом формате и с обязательной проверкой."
    )
    add_bullets(
        doc,
        "- выбрать ID и категорию (1..9) + краткое описание и проверку;\n"
        "- реализовать задачу в `ansible/roles/vuln_<scope>/tasks/vulns/<id>.yml`;\n"
        "- зарегистрировать ID в `ansible/roles/vuln_<scope>/defaults/main.yml` (список `vuln_<scope>_known`);\n"
        "- обновить `docs/guides/VULNERABILITY_CATALOG.md` (ID/категория/уровень/проверка);\n"
        "- обновить `docs/guides/VULNERABILITY_SCENARIOS.md` (сценарий: найти/доказать/исправить);\n"
        "- (опционально) включить ID в профили стенда `vuln_profiles.*`.",
    )
    doc.add_paragraph(
        "Изменение существующей уязвимости выполняется правкой соответствующего файла `tasks/vulns/<id>.yml` и обновлением "
        "методических материалов (проверки и сценария), чтобы требования к сдаче оставались однозначными."
    )

    doc.add_heading("2.8 Примеры реализации (несколько показательных)", level=2)
    doc.add_paragraph(
        "Ниже приведены несколько примеров, демонстрирующих принципы реализации для разных ОС и доменной среды. "
        "Полный перечень уязвимостей и сценариев проверки приведён в приложениях."
    )
    add_example_vuln_impl(
        doc,
        "Пример 1: Linux — sudo без пароля (linux.sudo.nopasswd)",
        "ansible/roles/vuln_linux/tasks/vulns/linux.sudo.nopasswd.yml",
        counters,
    )
    add_example_vuln_impl(
        doc,
        "Пример 2: Windows — включение LLMNR (windows.network.llmnr_enabled)",
        "ansible/roles/vuln_windows/tasks/vulns/windows.network.llmnr_enabled.yml",
        counters,
    )
    add_example_vuln_impl(
        doc,
        "Пример 3: AD/GPO — ослабление password policy домена (ad.gpo.password_policy_weak)",
        "ansible/roles/vuln_ad/tasks/vulns/ad.gpo.password_policy_weak.yml",
        counters,
    )

    doc.add_heading("2.9 Обязательная доменная среда (AD) и организационная автоматизация", level=2)
    doc.add_paragraph(
        "Для курса «Защита ОС» доменная инфраструктура используется как обязательный элемент: Domain Controller обеспечивает DNS домена "
        "и является точкой централизованного управления политиками (GPO). "
        "В проекте реализованы уязвимости AD/GPO (ad.* и ad.gpo.*), а также организационная автоматизация (OU/группы/линковка и фильтрация GPO), "
        "чтобы стенды были масштабируемыми при большом числе обучающихся."
    )
    add_bullets(
        doc,
        "- присоединение к домену реализовано ролью `ad_join` и включается профилем `ad_join_profiles.*`;\n"
        "- организационная часть (OU/группы/GPO scope) реализована ролью `ad_org` и выполняется на DC;\n"
        "- AD-специфичные уязвимости применяются ролью `vuln_ad` (в т.ч. через GPO).",
    )

    doc.add_heading("2.10 Требования к шаблонам и агентам (SSH, QEMU GA)", level=2)
    doc.add_paragraph(
        "Критически важно, чтобы шаблоны ВМ обеспечивали управляемость стенда: Linux управляется по SSH, Windows управляется по SSH "
        "(OpenSSH Server) с PowerShell, а QEMU Guest Agent обеспечивает корректные статусы и управляемые операции из Proxmox."
    )
    add_placeholder(doc, "screenshot", "Proxmox: включение QEMU Guest Agent (Options/Configuration)", counters)
    doc.add_page_break()

    # 3) Инструкции администратора
    doc.add_heading("3. Развертывание стендов и сопровождение (инструкции администратора)", level=1)
    doc.add_paragraph(
        "Ниже приведены требования и порядок действий для администратора, который разворачивает и поддерживает стенды."
    )
    doc.add_heading("3.1 Требования (ПО, Proxmox, сеть)", level=2)
    doc.add_paragraph("Свод требований вынесен в документ docs/guides/PREREQUISITES.md. Ниже — практический минимум:")
    # concise excerpt
    add_bullets(
        doc,
        "- Terraform 1.x\n"
        "- Ansible (2.15+ рекомендуется)\n"
        "- SSH доступ до ВМ (Linux и Windows по SSH)\n"
        "- Proxmox API token с правами на создание/клонирование ВМ\n"
        "- QEMU Guest Agent: включить agent=1 и установить агент внутри ВМ\n"
        "- Для домена: доступность DC по сети и DNS на DC для клиентов",
    )

    doc.add_heading("3.2 Файлы конфигурации (что заполнять)", level=2)
    doc.add_paragraph("Для корректной работы стенда администратор должен заполнить следующие файлы (по месту):")
    add_bullets(
        doc,
        "- Terraform: для каждой ВМ скопировать terraform.tfvars.example → terraform.tfvars и заполнить параметры Proxmox/сети;\n"
        "- Ansible accounts: group_vars/all/accounts.yml (локально; не коммитить если есть пароли/ключи);\n"
        "- Ansible AD: group_vars/all/ad.yml (domain join + организационные параметры OU/GPO scope);\n"
        "- Уязвимости: group_vars/all/vulnerabilities.yml (списки ID в vuln_profiles).",
    )
    doc.add_paragraph(
        "Особое внимание требует параметр ad_stand_computers: это список COMPUTERNAME доменных хостов, "
        "по которым роль организационной автоматизации переносит компьютеры в OU стенда и ограничивает применение GPO."
    )
    add_placeholder(doc, "listing", "Пример: stands/windows-stand/infrastructure/ansible/group_vars/all/ad.yml (заполненный)", counters)
    add_placeholder(doc, "listing", "Пример: stands/windows-stand/infrastructure/ansible/group_vars/all/accounts.yml (профили учёток)", counters)

    doc.add_heading("3.3 Подготовка Proxmox (один раз на площадку)", level=2)
    doc.add_paragraph(
        "Перед развертыванием стендов администратор подготавливает Proxmox VE: проверяет сеть (bridge), хранилища, доступ по API и права токена. "
        "Рекомендуется заранее завести сущности для управления доступом обучающихся: pools (по одному на стенд/группу)."
    )
    add_bullets(
        doc,
        "- создать/проверить bridge (например vmbr0), к которому будут подключаться ВМ;\n"
        "- создать/проверить storage для дисков и cloud-init (например local-lvm);\n"
        "- создать API Token (Datacenter → Permissions → API Tokens) и выдать права на уровне Datacenter/Node/Storage;\n"
        "- при необходимости подготовить pool-ы под группы/стенды (см. docs/guides/PROXMOX_STUDENT_ACCESS.md).",
    )
    add_placeholder(doc, "screenshot", "Proxmox: создание API Token и назначение прав", counters)
    add_placeholder(doc, "screenshot", "Proxmox: пример pool стенда и ACL на pool", counters)

    doc.add_heading("3.4 Подготовка шаблонов ВМ", level=2)
    doc.add_paragraph(
        "Шаблоны — ключевой элемент воспроизводимости. В проекте используются два подхода: "
        "Linux cloud-init шаблоны (включая UEFI/OVMF и cloud-init диск) и Windows шаблоны (подготовленные Packer) "
        "с включённым OpenSSH Server. Domain Controller считается заранее подготовленным (AD DS + домен создан)."
    )
    doc.add_heading("3.4.1 Linux cloud-init templates", level=3)
    add_bullets(
        doc,
        "- каталог: stands/linux-stand/infrastructure/templates/;\n"
        "- скрипты: PowerShell (.ps1) и bash (.sh) — для переносимости;\n"
        "- важные параметры: BIOS=OVMF, machine=q35, efidisk0, vga=std (для стабильной консоли noVNC);\n"
        "- в образе: включить SSH и (желательно) qemu-guest-agent.",
    )
    add_placeholder(doc, "listing", "Фрагмент скрипта создания Linux cloud-init template (UEFI/OVMF, cloud-init, vga)", counters)

    doc.add_heading("3.4.2 Windows templates через Packer", level=3)
    add_bullets(
        doc,
        "- каталог: stands/windows-stand/infrastructure/packer/;\n"
        "- подготовить variables.common.pkrvars.hcl и variables.secrets.pkrvars.hcl (из *.example);\n"
        "- в шаблоне обязательно: OpenSSH Server включён, учётка для Ansible существует, PowerShell доступен по SSH;\n"
        "- qemu-guest-agent/virtio-tools — рекомендуется.",
    )
    add_placeholder(doc, "listing", "Пример variables.common.pkrvars.hcl (без секретов)", counters)

    doc.add_heading("3.5 Развертывание стенда (Terraform + Ansible)", level=2)
    doc.add_paragraph(
        "Развертывание выполняется скриптами deploy.sh в каждом стенде. Скрипт обеспечивает последовательность: "
        "templates (опционально) → terraform apply → ansible-playbook. "
        "При переносе стенда IP адреса можно переопределить через переменные окружения."
    )
    add_placeholder(doc, "listing", "stands/windows-stand/infrastructure/scripts/deploy.sh (полный листинг)", counters)
    try:
        deploy_win = read_text("stands/windows-stand/infrastructure/scripts/deploy.sh")
        doc.add_paragraph("Фрагмент deploy.sh (Windows стенд):")
        add_code_block(doc, "\n".join(deploy_win.splitlines()[:55]))
    except Exception:
        pass

    doc.add_heading("3.6 Чек-лист администратора после деплоя", level=2)
    add_bullets(
        doc,
        "- ВМ созданы и запущены, диски/сеть присутствуют, cloud-init применён (Linux);\n"
        "- SSH доступен до Linux/Windows (Windows по SSH + PowerShell);\n"
        "- (если join включён) клиенты состоят в домене, DNS указывает на DC;\n"
        "- OU/группы/GPO применились (gpupdate /force на клиентах); \n"
        "- профили уязвимостей включены согласно group_vars/all/vulnerabilities.yml;\n"
        "- студенту выданы права в Proxmox только на pool стенда (Power + Console).",
    )
    add_placeholder(doc, "screenshot", "Проверка доменного членства Windows (System → Domain)", counters)
    add_placeholder(doc, "screenshot", "Проверка применения GPO (gpresult /r или rsop.msc)", counters)

    doc.add_heading("3.7 Типовые проблемы и диагностика (troubleshooting)", level=2)
    add_bullets(
        doc,
        "- VM не грузится (No bootable device): проверить, что в Terraform корректно созданы disk/cloudinit и boot order;\n"
        "- консоль noVNC сразу отключается: проверить VGA=std (serial0 может ломать консоль), UEFI/OVMF параметры;\n"
        "- Linux просит логин/пароль в консоли: cloud-init чаще настраивает вход по SSH-ключу (пароль может быть не задан);\n"
        "- Ansible не подключается к Windows: проверить OpenSSH Server, учетку, ansible_shell_type=powershell;\n"
        "- domain join не проходит: проверить доступность DC и DNS на DC; проверить учетку join и пароль.",
    )
    doc.add_page_break()

    doc.add_heading("3.8 Краткое резюме команд деплоя", level=2)
    add_bullets(
        doc,
        "- Windows стенд: stands/windows-stand/infrastructure/scripts/deploy.sh\n"
        "- Linux стенд: stands/linux-stand/infrastructure/scripts/deploy.sh",
    )
    doc.add_paragraph(
        "Windows-стенд разворачивается с обязательным DC. После domain join клиентов DC плейбук прогоняется повторно, "
        "чтобы переместить компьютеры в OU стенда и применить/ограничить GPO (организационная автоматизация)."
    )
    doc.add_heading("3.9 Внесение изменений и сопровождение", level=2)
    doc.add_paragraph(
        "Сопровождение стенда предполагает изменение конфигураций и повторное применение IaC. Типовые операции:"
    )
    add_bullets(
        doc,
        "- изменить сеть/ресурсы ВМ: правки Terraform в stands/*/infrastructure/terraform/* и повторный terraform apply;\n"
        "- изменить набор уязвимостей: правки group_vars/all/vulnerabilities.yml и повторный запуск ansible-playbook;\n"
        "- изменить доменную интеграцию/GPO scope: правки group_vars/all/ad.yml и повторный запуск плейбука DC;\n"
        "- добавить уязвимость: создать tasks/vulns/<id>.yml, добавить ID в *_known, обновить документы каталога/сценариев.",
    )

    doc.add_heading("3.10 Ограничение доступа обучающихся в Proxmox (RBAC)", level=2)
    doc.add_paragraph(
        "Модель доступа обучающегося основана на Proxmox RBAC: 1 стенд = 1 pool, права назначаются на pool. "
        "Обучающийся видит только ВМ своего стенда и может выполнять только операции Power/Console. "
        "Подробные шаги приведены в docs/guides/PROXMOX_STUDENT_ACCESS.md."
    )

    doc.add_heading("3.11 Пошаговое развертывание Windows-стенда (детально)", level=2)
    doc.add_paragraph(
        "Данный подраздел предназначен для администратора, который впервые разворачивает стенд. "
        "Он описывает последовательность действий, точки контроля и типовые ошибки. "
        "Рекомендуется выполнять шаги в указанном порядке (DC → клиенты → повторный прогон DC)."
    )
    add_placeholder(doc, "figure", "Схема выполнения деплоя Windows-стенда: Packer → Terraform → Ansible (DC → WS/Server → DC)", counters)
    doc.add_heading("3.11.1 Подготовка переменных Packer (опционально)", level=3)
    doc.add_paragraph(
        "Если администратор собирает Windows шаблоны самостоятельно, необходимо подготовить var-files Packer. "
        "Секреты (токены, пароли, ISO URLs) не коммитятся — они остаются только локально."
    )
    add_bullets(
        doc,
        "- скопировать variables.common.pkrvars.hcl.example → variables.common.pkrvars.hcl;\n"
        "- скопировать variables.secrets.pkrvars.hcl.example → variables.secrets.pkrvars.hcl;\n"
        "- указать node/storage/bridge, параметры ISO/VirtIO, учётные данные bootstrap (если требуется);\n"
        "- запустить сборку шаблонов (через build-example.sh или напрямую packer build).",
    )
    add_placeholder(doc, "screenshot", "Packer: успешное завершение build и появление шаблона в Proxmox", counters)
    doc.add_heading("3.11.2 Подготовка Terraform переменных (обязательно)", level=3)
    doc.add_paragraph(
        "Для каждой ВМ Windows-стенда (windows-10, windows-server, domain-controller) создаётся свой каталог Terraform "
        "и свой файл terraform.tfvars (локально). Это повышает прозрачность, так как разные ВМ могут иметь разный шаблон, IP и параметры."
    )
    add_placeholder(doc, "listing", "Пример terraform.tfvars для windows-10 (без секретов)", counters)
    add_code_block(
        doc,
        "\n".join(
            [
                "# Пример (без секретов): stands/windows-stand/infrastructure/terraform/windows-10/terraform.tfvars",
                "proxmox_api_url   = \"https://pve.example:8006/api2/json\"",
                "proxmox_node      = \"pve\"",
                "storage           = \"local-lvm\"",
                "proxmox_bridge    = \"vmbr0\"",
                "",
                "template_name     = \"windows-10-template\"",
                "windows_ws_ip     = \"192.168.101.10\"",
                "use_dhcp          = false",
                "cidr_prefix       = 24",
                "gateway           = \"192.168.101.1\"",
                "nameserver        = \"192.168.101.30\"  # DC как DNS",
                "",
                "# Операторский доступ (до развертывания):",
                "admin_user                = \"ansible\"",
                "bootstrap_admin_password   = \"\"  # хранить локально",
                "admin_password            = \"\"  # хранить локально",
                "admin_ssh_public_key      = \"ssh-ed25519 AAAA...\"",
            ]
        ),
    )
    doc.add_heading("3.11.3 Применение Terraform (создание ВМ)", level=3)
    doc.add_paragraph(
        "Terraform выполняется в каждом каталоге. После terraform apply важно проверить, что ВМ действительно стартовали "
        "и доступен SSH (на Windows — OpenSSH Server)."
    )
    add_bullets(
        doc,
        "- (cd stands/windows-stand/infrastructure/terraform/domain-controller) terraform init; terraform apply;\n"
        "- (cd stands/windows-stand/infrastructure/terraform/windows-10) terraform init; terraform apply;\n"
        "- (cd stands/windows-stand/infrastructure/terraform/windows-server) terraform init; terraform apply;",
    )
    add_placeholder(doc, "screenshot", "Proxmox: созданные ВМ (DC, Windows Server, Windows 10) и их состояние", counters)
    doc.add_heading("3.11.4 Запуск Ansible (DC → клиенты → DC)", level=3)
    doc.add_paragraph(
        "Ansible запускается из stands/windows-stand/infrastructure/ansible. "
        "Сначала настраивается Domain Controller (включая AD/GPO уязвимости и организационную часть), затем клиенты, "
        "после чего DC запускается повторно для пост-join действий (перемещение компьютеров в OU, линковка/фильтрация GPO)."
    )
    add_placeholder(doc, "listing", "Команды запуска Ansible плейбуков для Windows-стенда", counters)
    add_code_block(
        doc,
        "\n".join(
            [
                "# stands/windows-stand/infrastructure/ansible",
                "ansible-playbook domain-controller/playbook.yml",
                "ansible-playbook windows-10/playbook.yml",
                "ansible-playbook windows-server/playbook.yml",
                "ansible-playbook domain-controller/playbook.yml  # post-join (OU/GPO scope)",
            ]
        ),
    )
    doc.add_heading("3.11.5 Точки контроля (acceptance criteria)", level=3)
    add_bullets(
        doc,
        "- ВМ доступны по SSH: администратор может подключиться и выполнить команды проверки;\n"
        "- доменное членство клиентов корректно (если включено ad_join_profiles.*);\n"
        "- GPO применяются (gpupdate /force, gpresult /r);\n"
        "- уязвимости включены ровно те, что указаны в vulnerabilities.yml;\n"
        "- учётные записи/группы настроены согласно accounts.yml профилям.",
    )

    doc.add_heading("3.12 Пошаговое развертывание Linux-стенда (детально)", level=2)
    doc.add_paragraph(
        "Linux-стенд разворачивается из cloud-init шаблонов. Основное отличие — первичная конфигурация пользователя/SSH ключа выполняется cloud-init, "
        "поэтому важно обеспечить корректные параметры шаблона (UEFI/OVMF, cloud-init disk) и доступ по SSH."
    )
    add_placeholder(doc, "figure", "Схема Linux-стенда: Linux Server + Linux WS (+ опциональный domain join к DC)", counters)
    doc.add_heading("3.12.1 Сборка cloud-init шаблона", level=3)
    add_bullets(
        doc,
        "- скачать cloud-image (qcow2) подходящего дистрибутива;\n"
        "- создать шаблон через stands/linux-stand/infrastructure/templates;\n"
        "- проверить UEFI/OVMF и vga=std (для консоли), наличие cloud-init диска;\n"
        "- проверить, что sshd включён и доступ по ключу работает.",
    )
    doc.add_heading("3.12.2 Применение Terraform", level=3)
    doc.add_paragraph(
        "Terraform создаёт Linux Server и Linux WS. Сеть задаётся динамически: DHCP или статическая адресация по переменным. "
        "После создания ВМ проверяется доступность SSH и факт применения cloud-init параметров."
    )
    add_placeholder(doc, "listing", "Пример terraform.tfvars для Linux VM (без секретов)", counters)
    doc.add_heading("3.12.3 Ansible: учётки → (опц.) домен → уязвимости", level=3)
    doc.add_paragraph(
        "Плейбук Linux хоста выполняет роли в порядке: accounts → ad_join (если включено) → vuln_linux. "
        "Это гарантирует, что нужные пользователи существуют до применения уязвимостей и/или присоединения к домену."
    )
    add_placeholder(doc, "listing", "Команда запуска Linux плейбуков", counters)

    doc.add_heading("3.13 Управление секретами и требования к репозиторию", level=2)
    doc.add_paragraph(
        "Так как стенды содержат намеренно уязвимые настройки, важно отделять учебные уязвимости от секретов администратора. "
        "Секреты (tfvars, пароли join, токены API) не должны попадать в Git. "
        "В репозитории это обеспечивается .gitignore и использованием *.example файлов."
    )
    add_bullets(
        doc,
        "- хранить terraform.tfvars локально (не коммитить);\n"
        "- пароли домена и join креды хранить в отдельном защищённом хранилище;\n"
        "- при необходимости использовать Ansible Vault/менеджер секретов, но это не является обязательным для работы стенда;\n"
        "- регулярно проверять репозиторий на случайные секреты (pre-commit/сканеры).",
    )

    doc.add_heading("3.14 Изменение профилей уязвимостей (кратко)", level=2)
    doc.add_paragraph(
        "Оперативное управление набором уязвимостей выполняется без правок ролей: "
        "администратор редактирует только `vuln_profiles` в `group_vars/all/vulnerabilities.yml` и повторно запускает нужный playbook. "
        "Подробная модель (ID, роли, профили) описана в разделе 2.5–2.7."
    )
    add_placeholder(doc, "listing", "Пример: vulnerabilities.yml (профили ID уязвимостей)", counters)

    doc.add_heading("3.15 Расширение каталога уязвимостей (кратко)", level=2)
    doc.add_paragraph(
        "Добавление новой уязвимости выполняется по единому шаблону (ID → `tasks/vulns/<id>.yml` → регистрация в `*_known` → обновление каталога/сценариев). "
        "Подробная методика и примеры реализации приведены в разделе 2.7–2.8."
    )

    # 4) Сценарии тренировки
    doc.add_heading("4. Сценарии тренировки и обучения", level=1)
    doc.add_paragraph(
        "Сценарий обучения построен вокруг доступа обучающегося по VPN к Proxmox Web UI с ограниченными правами "
        "(только ВМ своего стенда). Обучающийся сам включает ВМ и выполняет задания по аудиту и hardening."
    )
    doc.add_heading("4.1 Роли и модель работы", level=2)
    add_bullets(
        doc,
        "- Администратор: разворачивает стенд, включает профили уязвимостей, выдаёт доступ обучающемуся;\n"
        "- Обучающийся: управляет ВМ через Proxmox, выполняет аудит/исправления, сдаёт отчёт;\n"
        "- Преподаватель: проверяет отчёт и корректность исправлений.",
    )
    doc.add_heading("4.2 Этапы занятия", level=2)
    doc.add_paragraph("Краткая структура сценария (см. docs/guides/TRAINING_SCENARIO.md):")
    add_bullets(
        doc,
        "- Блок A: проверка доменной среды (AD must-have)\n"
        "- Блок B: аудит базовой безопасности (учётки/политики, обновления/защита, логирование)\n"
        "- Блок C: исправление (hardening) и повторная проверка\n"
        "- Итог: отчёт с доказательствами «до/после»",
    )

    doc.add_page_break()
    doc.add_heading("Заключение", level=1)
    doc.add_paragraph(
        "В результате работы разработана платформа киберучений, обеспечивающая воспроизводимое развертывание учебной инфраструктуры "
        "в Proxmox VE и управляемое включение типовых уязвимостей/ошибок администрирования. "
        "Использование Terraform и Ansible позволяет стандартизировать процесс подготовки стендов, сократить ручные операции и "
        "обеспечить повторяемость лабораторных работ. Наличие доменной инфраструктуры (Active Directory) и набора GPO-уязвимостей "
        "приближает учебный процесс к реальным корпоративным условиям."
    )
    doc.add_paragraph(
        "Дальнейшее развитие может включать расширение каталога уязвимостей, автоматизацию сборки DC шаблона, а также добавление "
        "контрольных сценариев проверки (автотесты стенда) и интеграцию с CI для проверки корректности ролей/плейбуков."
    )

    doc.add_page_break()
    doc.add_heading("Приложения", level=1)
    doc.add_paragraph(
        "Приложения включают выдержки из эксплуатационной документации проекта. "
        "Раздел добавлен для удобства проверки и обеспечения требуемого объёма пояснительной записки."
    )
    add_placeholder(doc, "table", "Сводная таблица стендов (состав ВМ, IP, роли, профили уязвимостей)", counters)

    doc.add_heading("Приложение А — Требования и подготовка окружения", level=2)
    add_markdown_as_text(doc, prereq, counters)
    doc.add_page_break()

    doc.add_heading("Приложение Б — Сценарий тренировки/обучения", level=2)
    add_markdown_as_text(doc, training, counters)
    doc.add_page_break()

    doc.add_heading("Приложение В — Каталог уязвимостей (таблично)", level=2)
    add_markdown_as_text(doc, catalog_md, counters)
    doc.add_page_break()

    doc.add_heading("Приложение Г — Сценарии уязвимостей (подробно)", level=2)
    add_markdown_as_text(doc, scenarios_md, counters)
    doc.add_page_break()

    # Дополнительные приложения (эксплуатационные инструкции), чтобы не перегружать основной текст,
    # но сохранить полноту и требуемый объём ВКР.
    doc.add_heading("Приложение Д — Административный доступ (учётки, ключи, пароли)", level=2)
    add_markdown_as_text(doc, read_text("docs/guides/ADMIN_ACCESS.md"), counters)
    doc.add_page_break()

    doc.add_heading("Приложение Е — Управление пользователями и группами", level=2)
    add_markdown_as_text(doc, read_text("docs/guides/ACCOUNTS_AND_GROUPS.md"), counters)
    doc.add_page_break()

    doc.add_heading("Приложение Ж — Использование Ansible для включения уязвимостей", level=2)
    add_markdown_as_text(doc, read_text("docs/guides/ANSIBLE_VULNERABILITIES.md"), counters)
    doc.add_page_break()

    doc.add_heading("Приложение З — Доступ обучающихся к Proxmox (Pools/RBAC)", level=2)
    add_markdown_as_text(doc, read_text("docs/guides/PROXMOX_STUDENT_ACCESS.md"), counters)
    doc.add_page_break()

    doc.add_heading("Приложение И — Технологии проекта (обзор)", level=2)
    add_markdown_as_text(doc, read_text("docs/technical/TECHNOLOGIES.md"), counters)
    doc.add_page_break()

    doc.add_heading("Приложение К — Типовые ошибки безопасности (методические материалы)", level=2)
    add_markdown_as_text(doc, read_text("docs/technical/SECURITY_MISTAKES.md"), counters)
    doc.add_page_break()

    doc.add_heading("Приложение Л — Структура Ansible в проекте", level=2)
    add_markdown_as_text(doc, read_text("docs/structure/ansible-structure.md"), counters)
    doc.add_page_break()

    doc.add_heading("Приложение М — Структура Terraform (модули и конфигурации)", level=2)
    add_markdown_as_text(doc, read_text("docs/structure/terraform-modules-structure.md"), counters)
    doc.add_page_break()

    doc.add_heading("Приложение Н — Структура стендов", level=2)
    add_markdown_as_text(doc, read_text("stands/PROJECT_STRUCTURE.md"), counters)
    add_markdown_as_text(doc, read_text("stands/windows-stand/STRUCTURE.md"), counters)
    add_markdown_as_text(doc, read_text("stands/linux-stand/STRUCTURE.md"), counters)
    doc.add_page_break()

    doc.add_heading("Приложение О — Краткие инструкции по стендам (README)", level=2)
    add_markdown_as_text(doc, read_text("stands/README.md"), counters)
    add_markdown_as_text(doc, read_text("stands/windows-stand/README.md"), counters)
    add_markdown_as_text(doc, read_text("stands/linux-stand/README.md"), counters)
    doc.add_page_break()

    doc.add_heading("Список использованных источников", level=1)
    # Важно: ссылки в конце по требованию пользователя
    sources = [
        "NIST Special Publication 800-63B: Digital Identity Guidelines — Authentication and Lifecycle Management. `https://pages.nist.gov/800-63-3/sp800-63b.html` (дата обращения: "
        + date.today().isoformat()
        + ").",
        "MITRE ATT&CK® (база тактик и техник атак). `https://attack.mitre.org/` (дата обращения: "
        + date.today().isoformat()
        + ").",
        "CIS Benchmarks® (набор бенчмарков конфигурации). `https://www.cisecurity.org/cis-benchmarks` (дата обращения: "
        + date.today().isoformat()
        + ").",
        "Verizon Data Breach Investigations Report (DBIR) — страница отчёта. `https://www.verizon.com/business/resources/reports/dbir/` (дата обращения: "
        + date.today().isoformat()
        + ").",
        "Microsoft Digital Defense Report 2024 — страница отчёта. `https://www.microsoft.com/en-us/security/security-insider/threat-landscape/microsoft-digital-defense-report-2024` (дата обращения: "
        + date.today().isoformat()
        + ").",
        "Документация Terraform. `https://developer.hashicorp.com/terraform` (дата обращения: "
        + date.today().isoformat()
        + ").",
        "Документация Ansible. `https://docs.ansible.com/` (дата обращения: " + date.today().isoformat() + ").",
        "Документация Proxmox VE. `https://pve.proxmox.com/pve-docs/` (дата обращения: " + date.today().isoformat() + ").",
        "cloud-init documentation. `https://cloud-init.io/` (дата обращения: " + date.today().isoformat() + ").",
        "QEMU Guest Agent (обзор/документация). `https://www.qemu.org/docs/master/interop/qemu-ga.html` (дата обращения: "
        + date.today().isoformat()
        + ").",
        "OpenSSH в Windows (Microsoft Learn). `https://learn.microsoft.com/windows-server/administration/openssh/openssh_install_firstuse` (дата обращения: "
        + date.today().isoformat()
        + ").",
    ]
    for i, s in enumerate(sources, 1):
        doc.add_paragraph(f"{i}. {s}")

    out_path = save_docx_with_fallback(doc, target_path)
    wc = count_words_doc(doc)
    print(f"Updated: {out_path}")
    print(f"Word count: {wc} (rough pages ~ {estimate_pages(wc)})")


if __name__ == "__main__":
    main()

